# 20_make_ig_figures_and_tables.py
# Generate publication-ready figures + tables from IG outputs
# Inputs:
#   ig_one_demo.csv      (from 19_ig_explain_tabmfm.py)
#   ig_global_top.csv    (from 19_ig_explain_tabmfm.py)
# Outputs:
#   figures/ig_one_demo_top10.png
#   figures/ig_global_top15.png
#   tables/ig_one_demo_top10.csv
#   tables/ig_global_top15.csv
#   tables/ig_tables.docx

import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

IN_ONE = os.path.join(BASE_DIR, "ig_one_demo.csv")
IN_GLB = os.path.join(BASE_DIR, "ig_global_top.csv")

OUT_FIG_DIR = os.path.join(BASE_DIR, "figures")
OUT_TAB_DIR = os.path.join(BASE_DIR, "tables")
os.makedirs(OUT_FIG_DIR, exist_ok=True)
os.makedirs(OUT_TAB_DIR, exist_ok=True)

TOP_ONE = 10
TOP_GLB = 15

def safe_read_csv(path: str) -> pd.DataFrame:
    if not os.path.exists(path):
        raise FileNotFoundError(f"Not found: {path}\nPlease ensure you have generated it first.")
    return pd.read_csv(path)

def plot_barh(df: pd.DataFrame, feature_col: str, value_col: str, out_png: str, title: str):
    # Expect df already sorted descending by value
    dfp = df.copy()
    dfp = dfp.iloc[::-1]  # reverse for barh (largest on top)

    plt.figure(figsize=(8.5, 5.5))
    plt.barh(dfp[feature_col].astype(str), dfp[value_col].astype(float))
    plt.xlabel(value_col)
    plt.title(title)
    plt.tight_layout()
    plt.savefig(out_png, dpi=300)
    plt.close()

def to_docx_table(doc: Document, df: pd.DataFrame, title: str):
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for j, col in enumerate(df.columns):
        hdr_cells[j].text = str(col)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            v = row[col]
            if isinstance(v, float):
                cells[j].text = f"{v:.6f}"
            else:
                cells[j].text = str(v)
    doc.add_paragraph("")

def main():
    # --- Read ---
    one = safe_read_csv(IN_ONE)
    glb = safe_read_csv(IN_GLB)

    # --- Normalize columns (support both ig_norm/mean_ig_norm, etc.) ---
    # One-case
    if "ig_norm" not in one.columns:
        raise ValueError(f"{IN_ONE} must contain column 'ig_norm' (and 'feature'). Got columns: {list(one.columns)}")
    if "feature" not in one.columns:
        raise ValueError(f"{IN_ONE} must contain column 'feature'. Got columns: {list(one.columns)}")

    # Global
    if "mean_ig_norm" not in glb.columns:
        raise ValueError(f"{IN_GLB} must contain column 'mean_ig_norm' (and 'feature'). Got columns: {list(glb.columns)}")
    if "feature" not in glb.columns:
        raise ValueError(f"{IN_GLB} must contain column 'feature'. Got columns: {list(glb.columns)}")

    # --- Select top-k ---
    one_top = one.sort_values("ig_norm", ascending=False).head(TOP_ONE).reset_index(drop=True)
    glb_top = glb.sort_values("mean_ig_norm", ascending=False).head(TOP_GLB).reset_index(drop=True)

    # --- Save CSV tables (for paper) ---
    out_one_csv = os.path.join(OUT_TAB_DIR, f"ig_one_demo_top{TOP_ONE}.csv")
    out_glb_csv = os.path.join(OUT_TAB_DIR, f"ig_global_top{TOP_GLB}.csv")
    one_top.to_csv(out_one_csv, index=False, encoding="utf-8-sig")
    glb_top.to_csv(out_glb_csv, index=False, encoding="utf-8-sig")

    # --- Plots ---
    out_one_png = os.path.join(OUT_FIG_DIR, f"ig_one_demo_top{TOP_ONE}.png")
    out_glb_png = os.path.join(OUT_FIG_DIR, f"ig_global_top{TOP_GLB}.png")

    plot_barh(one_top, "feature", "ig_norm", out_one_png, f"Integrated Gradients (Single case) Top-{TOP_ONE}")
    plot_barh(glb_top, "feature", "mean_ig_norm", out_glb_png, f"Integrated Gradients (Global) Top-{TOP_GLB}")

    # --- Word doc with embedded tables + figures ---
    doc = Document()
    doc.add_heading("Integrated Gradients (IG) Summary Tables & Figures", level=1)

    doc.add_paragraph("This document is auto-generated from ig_one_demo.csv and ig_global_top.csv.")

    to_docx_table(doc, one_top, f"Table 1. Single-case IG Top-{TOP_ONE} (normalized)")
    doc.add_paragraph("Figure 1. Single-case IG (Top features).")
    doc.add_picture(out_one_png, width=Inches(6.5))
    doc.add_paragraph("")

    to_docx_table(doc, glb_top, f"Table 2. Global IG Top-{TOP_GLB} (mean normalized)")
    doc.add_paragraph("Figure 2. Global IG (Top features).")
    doc.add_picture(out_glb_png, width=Inches(6.5))
    doc.add_paragraph("")

    out_docx = os.path.join(OUT_TAB_DIR, "ig_tables.docx")
    doc.save(out_docx)

    print("=== Done ===")
    print("Saved tables:")
    print(" -", out_one_csv)
    print(" -", out_glb_csv)
    print("Saved figures:")
    print(" -", out_one_png)
    print(" -", out_glb_png)
    print("Saved Word:")
    print(" -", out_docx)

if __name__ == "__main__":
    main()
